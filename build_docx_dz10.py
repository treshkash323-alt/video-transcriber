# -*- coding: utf-8
"""Сборка DZ-10 docx: отчёт + ПЗ/РЭ. Скрины: screenshots/ — см. screenshots/README.md."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
SCREENSHOTS = ROOT / "screenshots"
GITHUB = "https://github.com/treshkash323-alt/video-transcriber"
GUIDE_MD = ROOT / "DZ-10_РУКОВОДСТВО_И_ПЗ.md"

FIGURES = [
    (
        "01-ui-3-success.png",
        "Рис. 1. Транскрибация — :8000, три задачи SUCCESS с текстом",
    ),
    (
        "02-flower-tasks.png",
        "Рис. 2. Flower — Workers, Succeeded: 3",
    ),
    (
        "03-flower-3-workers.png",
        "Рис. 3. Масштаб — три worker Online (docker compose --scale worker=3)",
    ),
    (
        "04-docker-compose.png",
        "Рис. 4. Запуск — docker compose up --build в терминале",
    ),
    (
        "05-docker-desktop.png",
        "Рис. 5. Docker Desktop — сервисы video-transcriber",
    ),
    (
        "06-history-after-restart.png",
        "Рис. 6. История SQLite — :8000 после docker compose down и повторного up",
    ),
    (
        "07-swagger-history.png",
        "Рис. 7. Swagger — GET /history возвращает список задач",
    ),
]

GUIDE_FIGURES = [
    ("02-flower-tasks.png", "Рис. A. Flower — Workers, Succeeded: 3"),
    ("02b-flower-tasks-list.png", "Рис. B. Flower — Tasks, transcribe_video SUCCESS"),
    ("03-flower-3-workers.png", "Рис. C. Flower — три worker Online (scale worker=3)"),
]


def setup_doc(title: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2)
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc


def meta(doc: Document, rows: list[tuple[str, str]]) -> None:
    for a, b in rows:
        p = doc.add_paragraph()
        p.add_run(a + " ").bold = True
        p.add_run(b)
    doc.add_paragraph()


def table(doc: Document, headers: list[str], rows: list[tuple]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def add_figure(doc: Document, filename: str, caption: str) -> str:
    path = SCREENSHOTS / filename
    p = doc.add_paragraph()
    p.add_run(caption).bold = True
    if path.is_file():
        doc.add_picture(str(path), width=Cm(15.5))
        status = "✅"
    else:
        run = doc.add_paragraph().add_run(f"[Вставить скрин: screenshots/{filename}]")
        run.italic = True
        status = "☐"
    doc.add_paragraph()
    return status


def _add_rich_paragraph(doc: Document, text: str, *, bullet: bool = False, quote: bool = False) -> None:
    text = text.strip()
    if not text:
        return
    style = "List Bullet" if bullet else "Normal"
    p = doc.add_paragraph(style=style)
    if quote:
        p.paragraph_format.left_indent = Cm(0.5)

    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).bold = True
        else:
            plain = re.sub(r"`([^`]+)`", r"\1", part)
            p.add_run(plain)


def _add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Cm(0.5)


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and "|" in s[1:-1]


def _parse_table(lines: list[str], idx: int) -> tuple[list[str], list[tuple], int]:
    header_line = lines[idx].strip()
    headers = [c.strip() for c in header_line.strip("|").split("|")]
    idx += 1
    if idx < len(lines) and re.match(r"^\|[-:\s|]+\|$", lines[idx].strip()):
        idx += 1
    rows: list[tuple] = []
    while idx < len(lines) and _is_table_row(lines[idx]):
        cells = [c.strip() for c in lines[idx].strip().strip("|").split("|")]
        while len(cells) < len(headers):
            cells.append("")
        rows.append(tuple(cells[: len(headers)]))
        idx += 1
    return headers, rows, idx


def md_to_docx(doc: Document, md_path: Path, *, start_heading: str | None = None) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    started = start_heading is None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not started:
            if stripped.startswith("## ") and start_heading in stripped:
                started = True
            else:
                i += 1
                continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            code: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            _add_code_block(doc, code)
            i += 1
            continue

        if _is_table_row(stripped):
            headers, rows, i = _parse_table(lines, i)
            table(doc, headers, rows)
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("> "):
            _add_rich_paragraph(doc, stripped[2:], quote=True)
            i += 1
            continue

        if stripped.startswith("- [") or stripped.startswith("- **"):
            _add_rich_paragraph(doc, stripped[2:], bullet=True)
            i += 1
            continue

        if stripped.startswith("- "):
            _add_rich_paragraph(doc, stripped[2:], bullet=True)
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            _add_rich_paragraph(doc, re.sub(r"^\d+\.\s", "", stripped), bullet=True)
            i += 1
            continue

        if stripped:
            _add_rich_paragraph(doc, stripped)

        i += 1


def _save_doc(doc: Document, primary: Path, fallback: Path) -> Path:
    try:
        doc.save(primary)
        return primary
    except PermissionError:
        doc.save(fallback)
        print("Файл открыт в Word — сохранено как:", fallback.name)
        return fallback


def build_report() -> Path:
    doc = setup_doc("ДЗ-10 — отчёт для сдачи (Lite)")
    meta(
        doc,
        [
            ("Студент:", "Игорь Кашинцев"),
            ("Курс:", "VibeCoder, Тема 10 — Celery, FastAPI"),
            ("Проект:", "video-transcriber"),
            ("GitHub:", GITHUB),
            ("Папка:", "Projects/ДЗ-10/video-transcriber/"),
            ("Запуск:", "localhost:8000 · Flower :5555"),
            ("Дата:", "09.06.2026 · готово к сдаче"),
        ],
    )

    doc.add_heading("1. Цель", level=1)
    doc.add_paragraph(
        "AI-транскрибатор: FastAPI принимает файлы, Celery + Whisper обрабатывает в фоне, "
        "Redis — очередь, Flower — мониторинг. Статусы PENDING → STARTED → SUCCESS. "
        "Масштаб: --scale worker=3. История задач — SQLite (data/history.db)."
    )

    doc.add_heading("2. Стек", level=1)
    doc.add_paragraph(
        "Docker Compose · FastAPI · Celery · Redis · Flower · OpenAI Whisper · "
        "SQLite (история) · FFmpeg"
    )

    doc.add_heading("3. Безопасность", level=1)
    table(
        doc,
        ["Проверка", "Результат"],
        [
            (".env в git", "Нет"),
            ("API key в коде", "Нет"),
            ("Лимит 25 МБ", "Да (UI + сервер)"),
            ("Redis наружу", "Нет"),
        ],
    )

    doc.add_heading("4. Скриншоты", level=1)
    doc.add_paragraph(
        "Скриншоты из папки screenshots/ (python build_docx_dz10.py — отчёт и ПЗ)."
    )

    figure_status: list[tuple[str, str, str]] = []
    for filename, caption in FIGURES:
        short = caption.split("—", 1)[0].strip()
        status = add_figure(doc, filename, caption)
        figure_status.append((filename.replace(".png", ""), short, status))

    doc.add_heading("5. Чеклист скринов", level=1)
    table(
        doc,
        ["Файл", "Описание", "Статус"],
        [(a, b, c) for a, b, c in figure_status],
    )

    doc.add_heading("6. Результат", level=1)
    table(
        doc,
        ["Пункт", "Статус"],
        [
            ("docker compose up — 4 сервиса", "✅"),
            ("3 файла транскрибированы", "✅"),
            ("Flower работает", "✅"),
            ("scale worker=3", "✅"),
            ("История SQLite", "✅"),
            ("GitHub без секретов", "✅"),
            ("Скрины в Word", "☐" if any(s == "☐" for *_, s in figure_status) else "✅"),
        ],
    )

    doc.add_heading("7. Комментарий для преподавателя", level=1)
    doc.add_paragraph(
        "Lite: FastAPI + Celery + Redis + Docker. Whisper через OpenAI API. "
        "Лимит 25 МБ учтён. Flower на том же образе, что worker. "
        "История задач в SQLite переживает docker compose down. "
        "Кэш по хешу — повторный файл без OpenAI. "
        "План развития: EDU + CRM — TODO_ROADMAP.md."
    )

    return _save_doc(
        doc,
        ROOT / "DZ-10_отчёт_для_сдачи.docx",
        ROOT / "DZ-10_отчёт_для_сдачи_новый.docx",
    )


def build_guide() -> Path:
    doc = setup_doc("ДЗ-10 — руководство по эксплуатации и пояснительная записка")
    meta(
        doc,
        [
            ("Студент:", "Игорь Кашинцев"),
            ("Курс:", "VibeCoder, Тема 10 — Celery, FastAPI"),
            ("Проект:", "video-transcriber — AI-транскрибатор видео"),
            ("GitHub:", GITHUB),
            ("Папка:", "Projects/ДЗ-10/video-transcriber/"),
            ("Запуск:", "localhost:8000 · Flower :5555"),
            ("Дата:", "09.06.2026 · готово к сдаче"),
        ],
    )

    md_to_docx(doc, GUIDE_MD, start_heading="1. Зачем этот проект")

    doc.add_heading("Приложение — скриншоты Flower и UI", level=1)
    doc.add_paragraph(
        "Ключевые скриншоты для разделов 6–8 (мониторинг Celery и масштабирование)."
    )
    for filename, caption in GUIDE_FIGURES:
        add_figure(doc, filename, caption)

    return _save_doc(
        doc,
        ROOT / "DZ-10_РУКОВОДСТВО_И_ПЗ.docx",
        ROOT / "DZ-10_РУКОВОДСТВО_И_ПЗ_новый.docx",
    )


if __name__ == "__main__":
    report_path = build_report()
    guide_path = build_guide()
    found = sum(1 for f, _ in FIGURES if (SCREENSHOTS / f).is_file())
    guide_found = sum(1 for f, _ in GUIDE_FIGURES if (SCREENSHOTS / f).is_file())
    print(f"OK отчёт: {report_path}")
    print(f"OK ПЗ/РЭ: {guide_path}")
    print(f"Скринов в отчёте: {found}/{len(FIGURES)}")
    print(f"Скринов в ПЗ: {guide_found}/{len(GUIDE_FIGURES)}")
